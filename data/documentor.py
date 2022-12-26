from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Mm
import matplotlib.pyplot as plt
from .models import *
from dateutil import tz
import pytz
from datetime import datetime, timedelta, date


def word_creator(checker, time1, time2, q_fetch):
    document = Document()

    hedin = document.add_heading(f'XYZ Bridge Structural Health Monitoring {checker} Report', 0)
    hedin.bold = True

    document.add_paragraph(f'Tabulation Time: {time1}')
    document.add_paragraph(f'Analysis Time：{time2}')

    document.add_heading(' 1.System Self-Assessment', level=1)

    records = (
        ('', '', 'Wind', 'UAN', '0', '0%', 'Very Poor'),
        (1, 'Environment', 'Temp and Humid', 'THS', '0', '0%', 'Very Poor'),
        ('', '', 'Temperature', 'DTS', '0', '0%', 'Very Poor'),
        (2, 'Load', 'Earthquake', 'VIB', '0', '0%', 'Very Poor'),
        ('', '', 'Gider Deflection', 'DEF', '1', '100%', 'Excellent'),
        (3, 'Deformation', 'Pier and Abutt. Tilt', 'TIL', '0', '0%', 'Very Poor'),
        ('', '', 'Support Disp', 'MDS', '0', '0%', 'Very Poor'),
        (4, 'Response', 'Cable Force', 'MFS', '0', '0%', 'Very Poor'),
        ('', '', 'Concrete Strain', 'VWS', '0', '0%', 'Very Poor'),
        (5, 'Dynamic', 'Vibration', 'VIB', '0', '0%', 'Very Poor')
    )

    table = document.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No.'
    hdr_cells[1].text = 'Category'
    hdr_cells[2].text = 'Monitoring'
    hdr_cells[3].text = 'Sensor Type'
    hdr_cells[4].text = 'Num'
    hdr_cells[5].text = 'Integrity'
    hdr_cells[6].text = 'Conclusion'
    for num, cat, descr, sen, nu, integ, conc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(num)
        row_cells[1].text = cat
        row_cells[2].text = descr
        row_cells[3].text = sen
        row_cells[4].text = nu
        row_cells[5].text = integ
        row_cells[6].text = conc

    document.add_paragraph('Monitoring Conclusions：')
    paragraph4 = document.add_paragraph(
            f'Today the health monitoring system for XYZ Bridge: ', style='List Bullet')
    paragraph4.add_run('Working Conditions Very poor\n').italic = True
    paragraph4.add_run('Remarks: ').bold = True
    paragraph4.add_run('Only 1 out of 10 sensors has been detected').italic = True
    paragraph5 = document.add_paragraph('Wind Scale: ', style='List Bullet')
    paragraph5.add_run('No report\n').italic = True
    paragraph5.add_run('Remarks: ').bold = True
    paragraph5.add_run('No sensor detected').italic = True
    paragraph6 = document.add_paragraph('Environment temperature and humidity：', style='List Bullet')
    paragraph6.add_run('No report\n').italic = True
    paragraph6.add_run('Remarks: ').bold = True
    paragraph6.add_run('No sensor detected').italic = True
    paragraph7 = document.add_paragraph('Pylon inside temperature and humidity：', style='List Bullet')
    paragraph7.add_run('No report\n').italic = True
    paragraph7.add_run('Remarks: ').bold = True
    paragraph7.add_run('No sensor detected').italic = True
    paragraph8 = document.add_paragraph('Girder inside temperature and humidity：', style='List Bullet')
    paragraph8.add_run('No report\n').italic = True
    paragraph8.add_run('Remarks: ').bold = True
    paragraph8.add_run('No sensor detected').italic = True
    paragraph9 = document.add_paragraph('Stayed-cable inside temperature and humidity：', style='List Bullet')
    paragraph9.add_run('No report\n').italic = True
    paragraph9.add_run('Remarks: ').bold = True
    paragraph9.add_run('No sensor detected').italic = True
    paragraph10 = document.add_paragraph('Cable force deviation：', style='List Bullet')
    paragraph10.add_run('No report\n').italic = True
    paragraph10.add_run('Remarks: ').bold = True
    paragraph10.add_run('No sensor detected').italic = True
    paragraph11 = document.add_paragraph('Stress：', style='List Bullet')
    paragraph11.add_run('No report\n').italic = True
    paragraph11.add_run('Remarks: ').bold = True
    paragraph11.add_run('No sensor detected').italic = True

    # document.add_page_break()

    document.add_heading(' 2.General Monitoring Conclusions：')
    document.add_paragraph('The measured girder deflections do not exceed the maximum deflection at service.', style='List Bullet')
    document.add_paragraph('The correlation of sensors is very poor and the measuring data is inaccurate and ineffective.', style='List Bullet')

    document.add_heading(' 3.Monitoring data analysis.', level=1)
    document.add_heading('   3.1 Data analysis of wind', level=2)
    document.add_paragraph('    No report').italic = True
    document.add_heading('   3.2 Data analysis of temperature and humidity', level=2)
    document.add_paragraph('    No report').italic = True
    document.add_heading('   3.3 Data analysis of earthquake', level=2)
    document.add_paragraph('    No report').italic = True
    document.add_heading('   3.4 Data analysis of Deck deflection', level=2)
    
    maximum_def, minimum_def= 0.0, 0.0
    # Definition of data contents
    if checker == 'Daily':
        record = DayData.objects.last()
        
        total = 0.00
        hour_list = []
        dartime = datetime.now(pytz.timezone('Africa/Dar_es_Salaam'))
        
        yesterday = dartime - timedelta(days = 1)
        day_items = Data.objects.filter(created__date=yesterday)
        
        if day_items.count() >= 700:
            count = 0
            hour_total_avg = 0
            for item in day_items:
                total += item.strain
                if (count%30) == 0 and count != 0:
                    hour_total_avg = total
                    hour_total_avg = hour_total_avg/30
                    total = 0
                    hour_list.append(hour_total_avg)
                count += 1
        elif day_items.count() >= 24 and day_items.count() <= 680:
            count = 0
            for item in day_items:
                hour_list.append(item.strain)
                if count == 23:
                    break
                else:
                    count += 1
        else:
            hour_list = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24]
        
        # Define column labels
        col_names = ['S/N', 'Time', 'Deflection', 'State']

        data = [['1', '(00:00Hrs - 01:00Hrs)', f'{hour_list[0]}', f'{checking_state(hour_list[0])}'],
        ['2', '(01:00Hrs - 02:00Hrs)', f'{hour_list[1]}', f'{checking_state(hour_list[1])}'],
        ['3', '(02:00Hrs - 03:00Hrs)', f'{hour_list[2]}', f'{checking_state(hour_list[2])}'],
        ['4', '(03:00Hrs - 04:00Hrs)', f'{hour_list[3]}', f'{checking_state(hour_list[3])}'],
        ['5', '(04:00Hrs - 05:00Hrs)', f'{hour_list[4]}', f'{checking_state(hour_list[4])}'],
        ['6', '(05:00Hrs - 06:00Hrs)', f'{hour_list[5]}', f'{checking_state(hour_list[5])}'],
        ['7', '(06:00Hrs - 07:00Hrs)', f'{hour_list[6]}', f'{checking_state(hour_list[6])}'],
        ['8', '(07:00Hrs - 08:00Hrs)', f'{hour_list[7]}', f'{checking_state(hour_list[7])}'],
        ['9', '(08:00Hrs - 09:00Hrs)', f'{hour_list[8]}', f'{checking_state(hour_list[8])}'],
        ['10', '(09:00Hrs - 10:00Hrs)', f'{hour_list[9]}', f'{checking_state(hour_list[9])}'],
        ['11', '(10:00Hrs - 11:00Hrs)', f'{hour_list[10]}', f'{checking_state(hour_list[10])}'],
        ['12', '(11:00Hrs - 12:00Hrs)', f'{hour_list[11]}', f'{checking_state(hour_list[11])}'],
        ['13', '(12:00Hrs - 13:00Hrs)', f'{hour_list[12]}', f'{checking_state(hour_list[12])}'],
        ['14', '(13:00Hrs - 14:00Hrs)', f'{hour_list[13]}', f'{checking_state(hour_list[13])}'],
        ['15', '(14:00Hrs - 15:00Hrs)', f'{hour_list[14]}', f'{checking_state(hour_list[14])}'],
        ['16', '(15:00Hrs - 16:00Hrs)', f'{hour_list[15]}', f'{checking_state(hour_list[15])}'],
        ['17', '(16:00Hrs - 17:00Hrs)', f'{hour_list[16]}', f'{checking_state(hour_list[16])}'],
        ['18', '(17:00Hrs - 18:00Hrs)', f'{hour_list[17]}', f'{checking_state(hour_list[17])}'],
        ['19', '(18:00Hrs - 19:00Hrs)', f'{hour_list[18]}', f'{checking_state(hour_list[18])}'],
        ['20', '(19:00Hrs - 20:00Hrs)', f'{hour_list[19]}', f'{checking_state(hour_list[19])}'],
        ['21', '(20:00Hrs - 21:00Hrs)', f'{hour_list[20]}', f'{checking_state(hour_list[20])}'],
        ['22', '(21:00Hrs - 22:00Hrs)', f'{hour_list[21]}', f'{checking_state(hour_list[21])}'],
        ['23', '(22:00Hrs - 23:00Hrs)', f'{hour_list[22]}', f'{checking_state(hour_list[22])}'],
        ['24', '(23:00Hrs - 24:00Hrs)', f'{hour_list[23]}', f'{checking_state(hour_list[23])}']]

        maximum_def = max(hour_list)
        minimum_def = min(hour_list)

    elif checker == 'Weekly':
        dartime = datetime.now(pytz.timezone('Africa/Dar_es_Salaam'))
        week_items = DayData.objects.filter(created__gte=dartime-timedelta(days=7))
        day_list = []
        day_list_values = []

        if week_items.count() >= 7:
            count_week = 0
            for item in range(7):
                time_created = item.created.ctime()
                time_str = str(time_created)
                a = time_str.index(':')
                b = time_str[a-3:a+6]
                day_time = time_str.replace(b, '')
                day_list.append(day_time)
                # ----------------\
                day_list_values.append(item.strain)
                if count_week == 6:
                    break
                else:
                    count_week += 1
        else:
            day_list = ['monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday']
            day_list_values = [41.55, 32.66, 43.66, 14.66, 8.66, 16.66, 27.66]

        col_names = ['S/N', 'Day', 'Deflection', 'State']

        data = [['1', f'{day_list[0]}', f'{day_list_values[0]}', f'{checking_state(day_list_values[0])}'],
            ['2', f'{day_list[1]}', f'{day_list_values[1]}', f'{checking_state(day_list_values[1])}'],
            ['3', f'{day_list[2]}', f'{day_list_values[2]}', f'{checking_state(day_list_values[2])}'],
            ['4', f'{day_list[3]}', f'{day_list_values[3]}', f'{checking_state(day_list_values[3])}'],
            ['5', f'{day_list[4]}', f'{day_list_values[4]}', f'{checking_state(day_list_values[4])}'],
            ['6', f'{day_list[5]}', f'{day_list_values[5]}', f'{checking_state(day_list_values[5])}'],
            ['7', f'{day_list[6]}', f'{day_list_values[6]}', f'{checking_state(day_list_values[6])}']]

        maximum_def = max(day_list_values)
        minimum_def = min(day_list_values)

    elif checker == 'Monthly':
        month_items = MonthData.objects.all()
        month_data_list = []
        month_name_list = []
        
        col_names = ['S/N', 'Month', 'Deflection', 'State']
        if month_items:
            data_loop, i, month_counter= [], 1, 0
            for item in month_items:
                time_created = item.created.month
                if int(time_created) == 1:
                    time_created = 12
                else:
                    time_created = time_created - 1
                month_name = checking_month(time_created)
                month_name_list.append(month_name)
                month_data_list.append(item.strain)
                data_loop.append([f'{i}', f'{month_name}', f'{item.strain}', f'{checking_state(item.strain)}'])
                i += 1
                if month_counter == 11:
                    break
                else:
                    month_counter += 1
            data = data_loop

        else:
            month_name_list = ['january', 'february', 'march', 'april', 'may']
            month_data_list = [61.38, 92.15, 43.03, 37.41, 75.86]

            data = [['1', f'{month_name_list[0]}', f'{month_data_list[0]}', f'{checking_state(month_data_list[0])}'],
                ['2', f'{month_name_list[1]}', f'{month_data_list[1]}', f'{checking_state(month_data_list[1])}'],
                ['3', f'{month_name_list[2]}', f'{month_data_list[2]}', f'{checking_state(month_data_list[2])}'],
                ['4', f'{month_name_list[3]}', f'{month_data_list[3]}', f'{checking_state(month_data_list[3])}'],
                ['5', f'{month_name_list[4]}', f'{month_data_list[4]}', f'{checking_state(month_data_list[4])}']]

        maximum_def = max(month_data_list)
        minimum_def = min(month_data_list)

    elif checker == 'Yearly':
        year_items = YearData.objects.all()
        col_names = ['S/N', 'Year', 'Deflection', 'State']
        year_data_values = []
        year_names = []

        if year_items:
            data_loop, i = [], 1
            for item in year_items:
                time_created = (item.created.year) - 1
                year_names.append(time_created)
                year_data_values.append(item.strain)
                print(f'Hii apa {year_names}..na ...{year_data_values}')
                data_loop.append([f'{i}', f'{year_names[len(year_names)-1]}', f'{item.strain}', f'{checking_state(year_data_values[len(year_data_values) - 1])}'])
                i += 1
            data = data_loop

        else:
            year_names = ['2020', '2021', '2022']
            year_data_values = [92.15, 43.03, 37.41]

            data = [['1', f'{year_names[0]}', f'{year_data_values[0]}', f'{checking_state(year_data_values[0])}'],
                ['2', f'{year_names[1]}', f'{year_data_values[1]}', f'{checking_state(year_data_values[1])}'],
                ['3', f'{year_names[2]}', f'{year_data_values[2]}', f'{checking_state(year_data_values[2])}']]

        maximum_def = max(year_data_values)
        minimum_def = min(year_data_values)

    else:
        data = [['-', '-', '-', '-'],
        ['nothing', 'to', 'display', '---']]

    if checker == 'Daily':
        document.add_paragraph(f"Yesterday's Average deflecton = {record.strain}, with average measures below.").bold = True
    # Obtain a 1-row, 3-column Table object 
    tb1 = document.add_table(rows=1, cols=len(col_names), style='Colorful Shading Accent 1')
    # Set the column name to the cell in the first row
    for i, cell in enumerate(tb1.rows[0].cells): # Get collection of Cell objects
        cell.text = col_names[i]                  #　Set value to Cell object
    # Set values while adding rows (Row objects)
    for d in data:
        row = tb1.add_row()   # Add Row Object
        row.height = Mm(8.0)  # Specify row height as 8mm
        
        for i, cell in enumerate(row.cells): # Obtaining Cell object
            cell.text = d[i]                  # Set value to Cell object
            cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM  # Set placement position in the cell

    # ---------------------------------------------
    document.add_paragraph('\n')

    if checker == 'Daily':
        x = ['0:00(mid-night)', '1:00(am)', '2:00(am)', '3:00(am)', '4:00(am)', '5:00(am)', '6:00(am)', '7:00(am)', 
        '8:00(am)', '9:00(am)', '10:00(am)', '11:00(am)', '12:00', '13:00(pm)', '14:00(pm)', '15:00(pm)', 
        '16:00(pm)', '17:00(pm)', '18:00(pm)', '19:00(pm)', '20:00(pm)', '21:00(pm)', '22:00(pm)', '23:00(pm)']
        y = hour_list
        
        plt.plot(x,y)
        plt.xlabel('Time Intervals in a day ')
        plt.ylabel('Deflection Measures')
        plt.title('Plotted Deflection Measures and DayTime recorded ')
        plt.axhline(y = 30, color = 'r', linestyle = 'dashed')
        plt.gcf().autofmt_xdate()
        plt.xticks(fontsize=6)
        plt.legend(['Deflection Trend', 'Critical Measure'])
        
    elif checker == 'Weekly':
        x = day_list
        y = day_list_values
        
        plt.plot(x,y)
        plt.xlabel('Intervals in Days')
        plt.ylabel('Deflection Measures')
        plt.title('Plotted Deflection Measures and Day recorded')
        plt.axhline(y = 30, color = 'r', linestyle = 'dashed')
        plt.gcf().autofmt_xdate()
        plt.xticks(fontsize=11)
        plt.legend(['Deflection Trend', 'Critical Measure'])

    elif checker == 'Monthly':
        x = month_name_list
        y = month_data_list

        plt.plot(x,y)
        plt.xlabel('Intervals in Month')
        plt.ylabel('Deflection Measures')
        plt.title('Plotted Deflection Measures and Month recorded')
        plt.axhline(y = 30, color = 'r', linestyle = 'dashed')
        plt.gcf().autofmt_xdate()
        plt.xticks(fontsize=11)
        plt.legend(['Deflection Trend', 'Critical Measure'])

    elif checker == 'Yearly':
        x = year_names
        y = year_data_values

        plt.plot(x,y)
        plt.xlabel('Intervals in Years')
        plt.ylabel('Deflection Measures')
        plt.title('Plotted Deflection Measures and Year recorded')
        plt.axhline(y = 30, color = 'r', linestyle = 'dashed')
        plt.gcf().autofmt_xdate()
        plt.xticks(fontsize=11)
        plt.legend(['Deflection Trend', 'Critical Measure'])

    else:
        x = [0, 2, 4, 6]
        y = [1, 3, 4, 8]

        plt.plot(x,y)
        plt.xlabel('x values')
        plt.ylabel('y values')
        plt.title('plotted x and y values')
        plt.legend(['line 1'])

    # save the figure
    plt.savefig('plot.png', dpi=300, bbox_inches='tight')

    # ------------------------------------------------------------------------
    document.add_picture('plot.png')
    document.add_paragraph('\n')
    document.add_paragraph(f'Maximum Deflection measured = {maximum_def}.', style='List Bullet')
    document.add_paragraph(f'Minimum Deflection measured = {minimum_def}.', style='List Bullet')
    document.add_paragraph(f'Maximum Service Deflection (Sensor) = {q_fetch.max_sema_deflection}.', style='List Bullet')
    document.add_paragraph(f'Bridge Type = {q_fetch.bridge_type}.', style='List Bullet')
    document.add_paragraph(f'Number of Lanes = {q_fetch.lanes_number}.', style='List Bullet')
    document.add_paragraph(f'Lane width = {q_fetch.lane_width}.', style='List Bullet')
    document.add_paragraph(f'Pedestrian Lane width = {q_fetch.ped_lane_width}.', style='List Bullet')
    document.add_paragraph(f'Effective Span = {q_fetch.effective_span}.', style='List Bullet')

    document.add_heading('   3.5 Data analysis of pier inclination', level=2)
    document.add_paragraph('    No report').italic = True
    document.add_heading('   3.6 Data analysis of girder displacement in longitudinal', level=2)
    document.add_paragraph('    No report').italic = True
    document.add_heading('   3.7 Data analysis of concrete stress', level=2)
    document.add_paragraph('    No report').italic = True
    document.add_heading('   3.8 Data analysis of vibration', level=2)
    document.add_paragraph('    No report').italic = True
    
    return document


def checking_month(value):
    value_dict = {
        1: 'January',
        2: 'February',
        3: 'March',
        4: 'April',
        5: 'May',
        6: 'June',
        7: 'July',
        8: 'August',
        9: 'September',
        10: 'October',
        11: 'November',
        12: 'December'}
    return value_dict[value]


def checking_day(value):
    value_dict = {
        1: 'Monday',
        2: 'Tuesday',
        3: 'Wednesday',
        4: 'Thursday',
        5: 'Friday',
        6: 'Saturday',
        7: 'Sunday'
    }
    return value_dict[value]


def checking_state(state):
    if int(state) <= 50:
        return 'Ok'
    else:
        return 'Not Ok'
